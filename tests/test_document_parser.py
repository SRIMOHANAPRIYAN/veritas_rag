"""Tests for document parsers."""

from pathlib import Path
from unittest.mock import patch, MagicMock
import pytest
import docx

from src.ingestion.document_parser import DocumentParser
from src.common.exceptions import IngestionError


@pytest.fixture
def parser():
    return DocumentParser()


def test_parse_txt(parser, tmp_path):
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("Hello world.\n\nThis is paragraph two.")

    blocks = parser.parse(txt_file)
    assert len(blocks) == 2
    assert blocks[0].text == "Hello world."
    assert blocks[1].text == "This is paragraph two."
    assert blocks[0].is_table is False


def test_parse_html(parser, tmp_path):
    html_file = tmp_path / "test.html"
    html_file.write_text(
        "<html><body>"
        "<h1>Main Title</h1>"
        "<p>Some text</p>"
        "<h2>Sub Title</h2>"
        "<table><tr><td>A</td><td>B</td></tr></table>"
        "</body></html>"
    )

    blocks = parser.parse(html_file)
    assert len(blocks) == 2
    assert blocks[0].text == "Some text"
    assert blocks[0].heading_path == ["Main Title"]
    assert blocks[1].heading_path == ["Main Title", "Sub Title"]
    assert blocks[1].is_table is True
    assert blocks[1].text == "A | B"


def test_parse_docx(parser, tmp_path):
    doc_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading("Doc Title", level=1)
    doc.add_paragraph("Docx paragraph.")
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    row[0].text = "C1"
    row[1].text = "C2"
    doc.save(doc_path)

    blocks = parser.parse(doc_path)
    assert len(blocks) == 3
    assert blocks[0].text == "Doc Title"
    assert blocks[1].text == "Docx paragraph."
    assert blocks[1].heading_path == ["Doc Title"]
    assert blocks[1].is_table is False
    assert blocks[2].is_table is True
    assert blocks[2].text == "C1 | C2"


def test_parse_missing_file(parser):
    with pytest.raises(IngestionError, match="File not found"):
        parser.parse(Path("does_not_exist.txt"))


def test_parse_unsupported_extension(parser, tmp_path):
    bad_file = tmp_path / "test.xyz"
    bad_file.write_text("data")
    with pytest.raises(IngestionError, match="Unsupported file extension"):
        parser.parse(bad_file)


def test_parse_size_limit(parser, tmp_path, monkeypatch):
    txt_file = tmp_path / "large.txt"
    txt_file.write_text("data")

    # Mock stat to return a large size
    original_stat = txt_file.stat

    class MockStat:
        def __init__(self, size):
            self.st_size = size

    def mock_stat(self, *args, **kwargs):
        if self.name == "large.txt":
            return MockStat(51 * 1024 * 1024)
        return original_stat(*args, **kwargs)

    monkeypatch.setattr(Path, "stat", mock_stat)

    with pytest.raises(IngestionError, match="exceeds the maximum allowed size"):
        parser.parse(txt_file)


@patch("src.ingestion.document_parser.pdfplumber")
def test_parse_pdf(mock_pdfplumber, parser, tmp_path):
    pdf_file = tmp_path / "test.pdf"
    pdf_file.write_text("dummy")  # create a file so exists() passes

    # Setup mock
    mock_pdf = MagicMock()
    mock_pdfplumber.open.return_value.__enter__.return_value = mock_pdf

    mock_page = MagicMock()
    mock_pdf.pages = [mock_page]

    # Mock tables
    mock_table = MagicMock()
    mock_table.bbox = (0, 0, 100, 100)
    mock_table.extract.return_value = [["Header1", "Header2"], ["Val1", "Val2"]]
    mock_page.find_tables.return_value = [mock_table]

    # Mock text outside tables
    mock_filtered_page = MagicMock()
    mock_filtered_page.extract_text.return_value = "PDF text."
    mock_page.filter.return_value = mock_filtered_page

    blocks = parser.parse(pdf_file)
    assert len(blocks) == 2
    assert blocks[0].is_table is True
    assert blocks[0].text == "Header1 | Header2\nVal1 | Val2"
    assert blocks[1].is_table is False
    assert blocks[1].text == "PDF text."
