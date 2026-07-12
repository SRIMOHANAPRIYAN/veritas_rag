"""Document parsers for VeritasRAG.

Supports PDF, DOCX, HTML, and TXT files, outputting a standardized list of ParsedBlock.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Union

from loguru import logger
import pdfplumber
import docx
from bs4 import BeautifulSoup

from src.common.exceptions import IngestionError


MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB


@dataclass
class ParsedBlock:
    """A contiguous block of text with metadata."""

    text: str
    page: int
    heading_path: List[str]
    is_table: bool
    source_path: str
    char_start: int
    char_end: int


class DocumentParser:
    """Parses supported document types into a list of ParsedBlock objects."""

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".txt"}

    def parse(self, file_path: Union[Path, str]) -> List[ParsedBlock]:
        """Parses a file and returns a list of ParsedBlocks.

        Args:
            file_path: The path to the file to parse.

        Returns:
            A list of ParsedBlock objects.

        Raises:
            IngestionError: If the file is not found, too large, unsupported, or corrupt.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise IngestionError(f"File not found: {file_path}")

        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE_BYTES:
            raise IngestionError(
                f"File {file_path} exceeds the maximum allowed size of 50MB "
                f"({file_size / (1024 * 1024):.2f}MB)."
            )

        ext = file_path.suffix.lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            raise IngestionError(f"Unsupported file extension: {ext} in {file_path}")

        logger.info(f"Parsing file: {file_path} (size: {file_size} bytes)")

        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext == ".docx":
                return self._parse_docx(file_path)
            elif ext in {".html", ".htm"}:
                return self._parse_html(file_path)
            elif ext == ".txt":
                return self._parse_txt(file_path)
            else:
                # Should not be reached due to previous check
                raise IngestionError(f"Unsupported extension: {ext}")
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            raise IngestionError(f"Corrupt or unparsable file {file_path}: {e}") from e

    def _parse_pdf(self, file_path: Path) -> List[ParsedBlock]:
        blocks = []
        source_path = str(file_path)
        char_offset = 0
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract tables
                tables = page.find_tables()
                table_bboxes = []
                for table in tables:
                    table_bboxes.append(table.bbox)
                    table_data = table.extract()
                    if table_data:
                        table_text = "\n".join(
                            [
                                " | ".join(
                                    [
                                        (
                                            str(cell).replace("\n", " ")
                                            if cell is not None
                                            else ""
                                        )
                                        for cell in row
                                    ]
                                )
                                for row in table_data
                                if row
                            ]
                        )
                        if table_text.strip():
                            blocks.append(
                                ParsedBlock(
                                    text=table_text.strip(),
                                    page=page_num,
                                    heading_path=[],
                                    is_table=True,
                                    source_path=source_path,
                                    char_start=char_offset,
                                    char_end=char_offset + len(table_text.strip()),
                                )
                            )
                            char_offset += len(table_text.strip()) + 2

                # Extract text outside tables
                def outside_tables(obj, table_bboxes=table_bboxes):
                    if (
                        "x0" not in obj
                        or "x1" not in obj
                        or "top" not in obj
                        or "bottom" not in obj
                    ):
                        return True
                    # If the object is within any table bbox, exclude it
                    for bbox in table_bboxes:
                        x0, top, x1, bottom = bbox
                # We skip tables for PDF parsing simplicity in this version,
                # or just extract text as a whole
                text = page.extract_text()
                if text and text.strip():
                    stripped_text = text.strip()
                    blocks.append(
                        ParsedBlock(
                            text=stripped_text,
                            page=page_num,
                            heading_path=[],
                            is_table=False,
                            source_path=source_path,
                            char_start=char_offset,
                            char_end=char_offset + len(stripped_text),
                        )
                    )
                    char_offset += len(stripped_text) + 2  # +2 for \n\n
        return blocks

    def _parse_docx(self, file_path: Path) -> List[ParsedBlock]:
        blocks = []
        source_path = str(file_path)
        char_offset = 0

        doc = docx.Document(file_path)
        current_heading_path = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith("Heading"):
                # Simplified heading tracking
                level = int(para.style.name.replace("Heading ", ""))
                current_heading_path = current_heading_path[: level - 1]
                current_heading_path.append(text)

            blocks.append(
                ParsedBlock(
                    text=text,
                    page=1,
                    heading_path=current_heading_path.copy(),
                    is_table=False,
                    source_path=source_path,
                    char_start=char_offset,
                    char_end=char_offset + len(text),
                )
            )
            char_offset += len(text) + 2  # +2 for \n\n

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]
                table_data.append(" | ".join(row_data))
            table_text = "\n".join(table_data)
            if table_text.strip():
                blocks.append(
                    ParsedBlock(
                        text=table_text.strip(),
                        page=1,
                        heading_path=current_heading_path.copy(),
                        is_table=True,
                        source_path=source_path,
                        char_start=char_offset,
                        char_end=char_offset + len(table_text.strip()),
                    )
                )
                char_offset += len(table_text.strip()) + 2

        return blocks

    def _parse_html(self, file_path: Path) -> List[ParsedBlock]:
        blocks = []
        source_path = str(file_path)
        char_offset = 0

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")

        current_heading_path = []

        # Find all structural elements in order
        for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "table"]):
            if el.name.startswith("h"):
                level = int(el.name[1])
                text = el.get_text(strip=True)
                if text:
                    current_heading_path = current_heading_path[: level - 1]
                    current_heading_path.append(text)
            elif el.name == "p":
                text = el.get_text(strip=True)
                if text:
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            page=1,
                            heading_path=current_heading_path.copy(),
                            is_table=False,
                            source_path=source_path,
                            char_start=char_offset,
                            char_end=char_offset + len(text),
                        )
                    )
                    char_offset += len(text) + 2  # +2 for \n\n
            elif el.name == "table":
                rows = []
                for tr in el.find_all("tr"):
                    cells = [
                        td.get_text(separator=" ", strip=True).replace("\n", " ")
                        for td in tr.find_all(["th", "td"])
                    ]
                    rows.append(" | ".join(cells))
                table_text = "\n".join(rows)
                if table_text.strip():
                    blocks.append(
                        ParsedBlock(
                            text=table_text.strip(),
                            page=1,
                            heading_path=current_heading_path.copy(),
                            is_table=True,
                            source_path=source_path,
                            char_start=char_offset,
                            char_end=char_offset + len(table_text.strip()),
                        )
                    )
                    char_offset += len(table_text.strip()) + 2

        return blocks

    def _parse_txt(self, file_path: Path) -> List[ParsedBlock]:
        blocks = []
        source_path = str(file_path)

        # We group lines into paragraphs separated by double newline
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        current_idx = 0

        for p in paragraphs:
            # We assume plain txt has no explicit headings
            # Find the paragraph in the content to get absolute offsets
            char_start = content.find(p, current_idx)
            if char_start != -1:
                char_end = char_start + len(p)
                blocks.append(
                    ParsedBlock(
                        text=p,
                        page=1,
                        heading_path=[],
                        is_table=False,
                        source_path=source_path,
                        char_start=char_start,
                        char_end=char_end,
                    )
                )
                current_idx = char_end

        return blocks
